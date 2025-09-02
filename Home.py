import streamlit as st
from docx import Document
import io
import datetime
from Office365_api import SharePoint  # Importação para integração com SharePoint
import json
import os
import uuid  # Adicionar esta linha

# Importações das novas funcionalidades
from navigation import criar_navegacao
from analytics import criar_dashboard_analytics
from text_analysis import criar_analise_texto
from controle_ideias import criar_sistema_controle
from gamificacao import criar_sistema_gamificacao
from notificacoes import criar_sistema_notificacoes
from cadastro_ideias import criar_formulario_ideia, listar_ideias
from mongodb_connection import mongo_manager
from auth import auth_manager  # Nova importação

# Configuração da página
st.set_page_config(
    page_title="Banco de Ideias - BIP",
    page_icon="ICON BIP.PNG",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'Get Help': 'https://www.extremelycoolapp.com/help',
        'Report a bug': "https://www.extremelycoolapp.com/bug",
        'About': "# BIP - Banco de Ideias e Práticas\nVersão 2.0 com funcionalidades avançadas!"
    }
)

# Função para limpar todos os campos com confirmação
def limpar_campos():
    # Resetar todos os campos do formulário
    st.session_state.anonimato_checkbox = False
    st.session_state.colaborador_select = None
    t.session_state.unidade_select = None
    st.session_state.categoria_select = None
    t.session_state.ideia_textarea = ""
    
    # Exibir mensagem de confirmação
    t.toast("Todos os campos foram limpos com sucesso!", icon="✅")

# Função para exibir diálogo de confirmação antes de limpar
def confirmar_limpeza():
    # Verificar se há dados preenchidos que seriam perdidos
    campos_preenchidos = [
        st.session_state.get("colaborador_select") is not None,
        st.session_state.get("unidade_select") is not None,
        st.session_state.get("categoria_select") is not None,
        st.session_state.get("ideia_textarea", "") != ""
    ]
    
# Função para criar e baixar documento
def criar_documento(anonimato, colaborador, unidade, categoria, ideia):
    # Criar um novo documento Word
    doc = Document()
    
    # Adicionar título
    doc.add_heading('Registro de Ideia - Banco de Ideias', 0)
    
    # Adicionar data e hora do registro
    data_hora = datetime.datetime.now().strftime("%d/%m/%Y %H:%M")
    doc.add_paragraph(f"Data de registro: {data_hora}")
    
    # Adicionar informações do colaborador (se não for anônimo)
    if not anonimato and colaborador:
        doc.add_paragraph(f"Colaborador: {colaborador}")
        doc.add_paragraph(f"Unidade: {unidade}")
    else:
        doc.add_paragraph("Colaborador: Anônimo")
    
    # Adicionar categoria e ideia
    doc.add_paragraph(f"Categoria: {categoria}")
    
    # Adicionar a ideia com formatação melhorada
    p = doc.add_paragraph()
    p.add_run("Ideia:").bold = True
    doc.add_paragraph(ideia)
    
    # Salvar o documento em um buffer de bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes

# Função para verificar a conexão com o SharePoint
def verificar_conexao_sharepoint():
    try:
        sharepoint = SharePoint()
        # Tenta autenticar
        conn = sharepoint._auth()
        return True
    except Exception as e:
        st.sidebar.warning(f"Não foi possível conectar ao SharePoint. Os arquivos serão salvos apenas localmente.")
        return False

# Inicializa os estados da sessão se não existirem
def inicializar_sessao():
    # Estados existentes
    if 'anonimato_checkbox' not in st.session_state:
        st.session_state.anonimato_checkbox = False
    if 'colaborador_select' not in st.session_state:
        st.session_state.colaborador_select = None
    if 'unidade_select' not in st.session_state:
        st.session_state.unidade_select = None
    if 'categoria_select' not in st.session_state:
        st.session_state.categoria_select = None
    if 'ideia_textarea' not in st.session_state:
        st.session_state.ideia_textarea = ""
    
    # Novos estados para as funcionalidades
    if 'usuario_pontos' not in st.session_state:
        st.session_state.usuario_pontos = 0
    if 'usuario_badges' not in st.session_state:
        st.session_state.usuario_badges = []
    if 'ideias_enviadas' not in st.session_state:
        st.session_state.ideias_enviadas = []
    if 'notificacoes' not in st.session_state:
        st.session_state.notificacoes = []
    
    # Verificar conexão com SharePoint
    st.session_state.sharepoint_conectado = verificar_conexao_sharepoint()

# Função principal da interface
def main():
    inicializar_sessao()
    
    # Sistema de navegação
    pagina_selecionada = criar_navegacao()
    
    # Verificar autenticação para páginas restritas
    if not auth_manager.require_auth(pagina_selecionada):
        return  # Para a execução se não autenticado
    
    # Roteamento baseado na página selecionada
    if pagina_selecionada == "🏠 Enviar Ideia":
        # Código original do formulário (acesso livre)
        st.header("BANCO DE IDEIAS e BOAS PRÁTICAS - REDE LIUS", divider="orange")
        st.write("""Ferramenta de registro e acompanhamento de ideias e boas práticas institucionais da Rede Lius""")
        
        # Sidebar com opções de identificação
        with st.sidebar:
            criar_sidebar()
        
        st.write("---")
        
        # Área principal para entrada da ideia
        criar_formulario_ideia()
    
    elif pagina_selecionada == "📊 Dashboard":
        st.header("📊 Dashboard Analytics", divider="blue")
        criar_dashboard_analytics()
    
    elif pagina_selecionada == "☁️ Análise de Texto":
        st.header("☁️ Análise de Texto", divider="green")
        criar_analise_texto()
    
    elif pagina_selecionada == "📋 Controle de Ideias":
        st.header("📋 Controle de Ideias", divider="red")
        criar_sistema_controle()
    
    elif pagina_selecionada == "🎮 Gamificação":
        st.header("🎮 Sistema de Gamificação", divider="violet")
        criar_sistema_gamificacao()
    
    elif pagina_selecionada == "🔔 Notificações":
        st.header("🔔 Sistema de Notificações", divider="orange")
        criar_sistema_notificacoes()
    
    elif pagina_selecionada == "🤖 Análise IA":
        st.header("🤖 Análise com IA", divider="rainbow")
        from ia_analysis import criar_analise_ia
        criar_analise_ia()
    
    elif pagina_selecionada == "📋 Listar Ideias":
        st.header("📋 Lista de Ideias", divider="blue")
        listar_ideias()
        
        st.write("---")
        st.write("Status da conexão:")
        
        # Exibir status da conexão com SharePoint
        if st.session_state.get('sharepoint_conectado', False):
            st.sidebar.success("✅ Conexão com o SharePoint feita com sucesso")
        else:
            st.sidebar.warning("⚠️ Sem conexão com SharePoint")

# Função para criar a sidebar
def criar_sidebar():
    
    st.header("Opções")
    
    # Caixa de seleção para anonimato
    anonimato = st.checkbox(
        "NÃO quero me identificar",
        value=st.session_state.get("anonimato_checkbox", False),
        key="anonimato_checkbox",
        help="Marque esta opção se deseja enviar sua ideia anonimamente",
        on_change=lambda: st.session_state.update(anonimato_checkbox=st.session_state.anonimato_checkbox)
    )
    
    # Nome do colaborador. Caso o checkbox "NÃO quero me identificar" esteja marcado, desabilitar este campo
    colaboradores = st.text_input(
        "Informe seu nome completo", 
        value=st.session_state.get("colaboradores", ""),
        key="colaboradores",
        disabled=st.session_state.anonimato_checkbox
    )
    
    colaborador = colaboradores
    
    # Unidades disponíveis
    unidades = [
        "CSA - BH",
        "CSA - CTG",
        "CSA - NL",
        "CSA - GZ",
        "CSA - DV",
        "EPSA",
        "ESA",
        "AIACOM",
        "ADEODATO",
        "SIC - SEDE",
        "PROVÍNCIA AGOSTINIANA"
    ]
    
    # Selectbox para escolha da unidade
    unidade = st.selectbox(
        "Selecione sua unidade:",
        unidades,
        index=None,
        key="unidade_select"
    )
    
    # Mensagem informativa sobre anonimato
    if anonimato:
        st.info("Fique tranquilo! Seu nome não será registrado no cadastro da ideia!")
    else:
        st.info("Seu nome e sua unidade serão registrados no cadastro da ideia!")
    
    # Botão para limpar campos com estilo e ícone
    #st.button(
        #"🗑️ Limpar campos", 
        #on_click=confirmar_limpeza, 
       # key='limpar_campos',
        #help="Clique para limpar todos os campos do formulário"
    #)

    st.write("---")

    st.write("Status da conexão:")


    # Exibir status da conexão com SharePoint
    if st.session_state.get('sharepoint_conectado', False):
        st.sidebar.success("✅ Conexão com o SharePoint feita com sucesso")
    else:
        st.sidebar.warning("⚠️ Sem conexão com SharePoint")

# Função para criar o formulário principal
def criar_formulario_ideia():
    # Categorias disponíveis
    categorias = [
        "Tecnologia & Inovação",
        "Currículo & Metodologia",
        "Infraestrutura & Espaço Físico",
        "Bem Estar & Cultura Escolar",
        "Eventos & Engajamento",
        "Gestão eficiênte de Custos",
        "Retenção de alunos",
        "Captação de alunos",
        "Rede Sociais"
    ]
    
    # Selectbox para escolha da categoria
    categoria = st.selectbox(
        "Sua ideia envolve que tipo de categoria?:",
        categorias,
        index=None,
        placeholder="Escolha uma das opções a seguir",
        key="categoria_select"
    )
    
    # Área de texto para a ideia
    ideia = st.text_area(
        "Escreva aqui a sua ideia",
        placeholder="Explique sua proposta em detalhes: o que é, como funcionaria, qual problema ou oportunidade atende.",
        key="ideia_textarea",
        height=200  # Altura ajustável para melhor visualização
    )
    
    # Contador de caracteres com formatação condicional
    caracteres = len(ideia)
    if caracteres > 0:
        if caracteres < 50:
            st.warning(f"Você escreveu apenas {caracteres} caracteres. Recomendamos detalhar mais sua ideia.")
        else:
            st.write(f"Você escreveu {caracteres} caracteres")
    
    # Botão para salvar a ideia
    if st.button("Enviar ideia", type="primary", key="Salvar_ideia"):
        processar_salvamento()

# Função para processar o salvamento da ideia
def processar_salvamento():
    # Obter valores atuais dos campos
    anonimato = st.session_state.anonimato_checkbox
    colaborador = st.session_state.colaboradores
    unidade = st.session_state.unidade_select
    categoria = st.session_state.categoria_select
    ideia = st.session_state.ideia_textarea
    
    # Validar campos obrigatórios
    if not categoria:
        st.error("Por favor, selecione uma categoria antes de salvar.")
        return
    
    if not ideia:
        st.error("Por favor, escreva sua ideia antes de salvar.")
        return
    
    if not anonimato and not colaborador:
        st.error("Por favor, selecione seu nome ou marque a opção de anonimato.")
        return
    
    if not unidade:
        st.error("Por favor, selecione sua unidade antes de salvar.")
        return
    
    # Criar documento
    doc_bytes = criar_documento(anonimato, colaborador, unidade, categoria, ideia)
    
    # Criar nome do arquivo (usando a categoria, removendo caracteres inválidos)
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M")
    filename = f"Ideia_{categoria.replace(' ', '_').replace('&', 'e')}_{timestamp}.docx"
    
    # Preparar dados para MongoDB
    import uuid
    ideia_data = {
        "id_unico": str(uuid.uuid4()),
        "titulo": f"Ideia - {categoria}",
        "autor": "Anônimo" if anonimato else colaborador,
        "email": "",  # Pode ser adicionado ao formulário se necessário
        "categoria": categoria,
        "unidade": unidade,
        "prioridade": "Média",  # Valor padrão
        "impacto": "Médio",     # Valor padrão
        "descricao": ideia,
        "justificativa": "",
        "recursos": "",
        "beneficios": "",
        "prazo_implementacao": "3-6 meses",  # Valor padrão
        "orcamento_estimado": "R$ 1.000 - R$ 5.000",  # Valor padrão
        "tags": [categoria.lower().replace(' ', '_')],
        "status": "Pendente",
        "votos": 0,
        "comentarios": [],
        "data_submissao": datetime.datetime.now().isoformat(),
        "arquivo_sharepoint": filename,
        "anonimo": anonimato
    }
    
    # Variáveis para controlar o sucesso das operações
    sharepoint_sucesso = False
    mongodb_sucesso = False
    
    # 1. SALVAR NO MONGODB
    with st.spinner("Salvando ideia no MongoDB..."):
        try:
            ideia_id = mongo_manager.salvar_ideia(ideia_data)
            if ideia_id:
                mongodb_sucesso = True
                st.success(f"✅ Ideia salva no MongoDB! ID: {ideia_id}")
            else:
                st.error("❌ Erro ao salvar no MongoDB")
        except Exception as e:
            st.error(f"❌ Erro ao salvar no MongoDB: {str(e)}")
    
    # 2. SALVAR NO SHAREPOINT
    with st.spinner("Salvando ideia no SharePoint..."):
        try:
            # Definir a pasta do SharePoint onde o arquivo será salvo
            sharepoint_folder = "Banco_de_Ideias"  # Pasta no SharePoint para armazenar as ideias
            
            # Fazer upload do arquivo para o SharePoint
            sharepoint = SharePoint()
            response = sharepoint.upload_file(filename, sharepoint_folder, doc_bytes.getvalue())
            
            # Verificar se o upload foi bem-sucedido
            if response:
                sharepoint_sucesso = True
                st.success(f"✅ Ideia salva no SharePoint!")
            else:
                st.warning("⚠️ Não foi possível salvar no SharePoint")
        except Exception as e:
            st.error(f"❌ Erro ao salvar no SharePoint: {str(e)}")
    
    # 3. RESUMO DO SALVAMENTO
    st.write("---")
    st.subheader("📋 Resumo do Salvamento")
    
    col1, col2 = st.columns(2)
    with col1:
        if mongodb_sucesso:
            st.success("✅ MongoDB: Salvo com sucesso")
        else:
            st.error("❌ MongoDB: Falha no salvamento")
    
    with col2:
        if sharepoint_sucesso:
            st.success("✅ SharePoint: Salvo com sucesso")
        else:
            st.error("❌ SharePoint: Falha no salvamento")
    
    # 4. OFERECER DOWNLOAD LOCAL
    st.download_button(
        label="📥 Baixar Ideia em Word",
        data=doc_bytes,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="download_ideia"
    )
    
    # 5. MENSAGEM FINAL
    if mongodb_sucesso and sharepoint_sucesso:
        st.success("🎉 Ideia salva com sucesso em ambos os sistemas!")
        st.balloons()
    elif mongodb_sucesso or sharepoint_sucesso:
        st.warning("⚠️ Ideia salva parcialmente. Verifique os detalhes acima.")
    else:
        st.error("❌ Falha ao salvar a ideia. Você ainda pode baixar o arquivo localmente.")

# Executar o aplicativo
if __name__ == "__main__":
    main()


def criar_relatorios():
    st.header("📈 Relatórios")
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("📊 Relatório Mensal"):
            st.success("Relatório mensal gerado!")
    
    with col2:
        if st.button("📈 Relatório de Tendências"):
            st.success("Relatório de tendências gerado!")

def criar_configuracoes():
    st.header("⚙️ Configurações")
    
    st.subheader("Configurações Gerais")
    
    # Configurações de notificação
    st.checkbox("Ativar notificações por email", value=True)
    st.checkbox("Ativar gamificação", value=True)
    
    # Configurações de análise
    st.selectbox("Idioma para análise de texto", ["Português", "Inglês", "Espanhol"])
    
    # Configurações de SharePoint
    st.text_input("URL do SharePoint", value="https://...")
    
    if st.button("💾 Salvar Configurações"):
        st.success("Configurações salvas com sucesso!")

# Rodapé com copyright
st.sidebar.markdown("""
    <style>
    .footer {
        position: fixed;
        left: 0;
        bottom: 0;
        width: 100%;
        background-color: #f0f0f0;
        color: #333;
        text-align: center;
        padding: 10px;
        font-size: 14px;
    }
    </style>
    <div class="footer">
        © 2025 FP&A e Orçamento - Rede Lius. Todos os direitos reservados.
    </div>
    """, unsafe_allow_html=True)





# REMOVER todo código duplicado e solto após esta linha